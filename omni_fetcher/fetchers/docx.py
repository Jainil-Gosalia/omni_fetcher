"""DOCX fetcher for OmniFetcher."""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import Any

import httpx

try:
    import docx
except ImportError:
    raise ImportError(
        "python-docx is required for DOCX fetching. Install with: pip install omni_fetcher[office]"
    )

from omni_fetcher.core.registry import source
from omni_fetcher.fetchers.base import BaseFetcher
from omni_fetcher.schemas.atomics import (
    ImageDocument,
    SheetData,
    SpreadsheetDocument,
    TextDocument,
    TextFormat,
)
from omni_fetcher.schemas.base import DataCategory, FetchMetadata, MediaType
from omni_fetcher.schemas.documents import DOCXDocument


@source(
    name="docx",
    uri_patterns=["*.docx"],
    mime_types=["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
    priority=20,
    description="Fetch and parse Microsoft Word documents",
)
class DOCXFetcher(BaseFetcher):
    """Fetcher for Microsoft Word documents."""

    name = "docx"
    priority = 20

    def __init__(self, timeout: float = 60.0):
        super().__init__()
        self.timeout = timeout

    @classmethod
    def can_handle(cls, uri: str) -> bool:
        """Check if this is a DOCX URL."""
        return uri.lower().endswith(".docx")

    async def fetch(self, uri: str, **kwargs: Any) -> DOCXDocument:
        """Fetch and parse a DOCX document."""
        if uri.startswith("file://"):
            docx_data = await self._fetch_local(uri)
        else:
            docx_data = await self._fetch_remote(uri)

        metadata = FetchMetadata(
            source_uri=uri,
            fetched_at=datetime.now(),
            source_name=self.name,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_size=len(docx_data),
            fetch_duration_ms=None,
            cache_hit=False,
            status_code=None,
            encoding=None,
            last_modified=None,
            etag=None,
        )

        docx_info = await self._parse_docx(docx_data, uri)

        text_doc = docx_info.get("text")
        if text_doc is None:
            text_doc = TextDocument(source_uri=uri, content="", format=TextFormat.PLAIN)

        return DOCXDocument(
            metadata=metadata,
            text=text_doc,
            images=docx_info.get("images", []),
            tables=docx_info.get("tables", []),
            page_count=docx_info.get("page_count"),
            author=docx_info.get("author"),
            title=docx_info.get("title"),
            created_at=docx_info.get("created_at"),
            modified_at=docx_info.get("modified_at"),
            category=DataCategory.DOCUMENT,
            media_type=MediaType.APPLICATION_DOCX,
        )

    async def _fetch_local(self, uri: str) -> bytes:
        """Fetch local DOCX file."""
        path = uri.replace("file://", "")
        if path.startswith("/"):
            path = path[1:]
        return Path(path).read_bytes()

    async def _fetch_remote(self, uri: str) -> bytes:
        """Fetch remote DOCX file."""
        headers = self.get_auth_headers()
        async with httpx.AsyncClient(timeout=self.timeout) as client:
            response = await client.get(uri, headers=headers)
            response.raise_for_status()
            return response.content

    async def _parse_docx(self, docx_data: bytes, uri: str) -> dict[str, Any]:
        """Parse DOCX and extract information."""

        def _parse():
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            doc = Document(io.BytesIO(docx_data))

            info: dict[str, Any] = {
                "text": None,
                "images": [],
                "tables": [],
                "page_count": None,
                "author": None,
                "title": None,
                "created_at": None,
                "modified_at": None,
            }

            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            full_text = "\n".join(text_parts)

            info["text"] = TextDocument(
                source_uri=uri,
                content=full_text,
                format=TextFormat.PLAIN,
            )

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                if table_data:
                    sheet_data = SheetData(
                        name=f"Table_{len(info['tables']) + 1}",
                        headers=table_data[0] if table_data else None,
                        rows=table_data[1:] if len(table_data) > 1 else [],
                        row_count=len(table_data),
                        col_count=len(table_data[0]) if table_data else 0,
                    )
                    info["tables"].append(
                        SpreadsheetDocument(
                            source_uri=uri,
                            sheets=[sheet_data],
                            format="docx",
                            sheet_count=1,
                        )
                    )

            core_props = doc.core_properties
            info["title"] = core_props.title or None
            info["author"] = core_props.author or None
            info["created_at"] = core_props.created or None
            info["modified_at"] = core_props.modified or None

            for rel in doc.part.rels.values():
                if rel.reltype == RT.IMAGE:
                    image_part = rel.target_part

                    info["images"].append(
                        ImageDocument(
                            source_uri=uri,
                            format=image_part.content_type,
                            raw=image_part.blob,
                            file_name=rel.target_ref.split("/")[-1]
                            if "/" in rel.target_ref
                            else rel.target_ref,
                        )
                    )

            return info

        import asyncio

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _parse)
